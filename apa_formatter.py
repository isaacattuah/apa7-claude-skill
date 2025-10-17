#!/usr/bin/env python3
"""
APA 7th Edition Document Formatter
Converts raw text/markdown to properly formatted APA documents
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re


class APAFormatter:
    """Format documents according to APA 7th Edition guidelines."""
    
    def __init__(self):
        self.doc = None
        
    def create_document(self, title_data, body_content, references):
        """
        Create an APA 7th edition formatted document.
        
        Parameters:
        -----------
        title_data : dict
            {'title': str, 'author': str, 'institution': str, 
             'course': str, 'instructor': str, 'date': str}
        body_content : list of dict
            [{'type': 'heading', 'level': int, 'text': str},
             {'type': 'paragraph', 'text': str},
             {'type': 'math', 'formula': str}]
        references : list of str
            ['Author, A. A. (Year). Title...', ...]
        
        Returns:
        --------
        Document object ready to save
        """
        self.doc = Document()
        
        # Configure document defaults
        self._set_document_defaults()
        
        # Set margins
        self._set_margins()
        
        # Add page numbers
        self._add_page_numbers()
        
        # Create title page
        self._create_title_page(title_data)
        
        # Add page break before body
        self.doc.add_page_break()
        
        # Add body content
        self._create_body(body_content)
        
        # Add references on new page
        self.doc.add_page_break()
        self._create_references(references)
        
        return self.doc
    
    def _set_document_defaults(self):
        """Set default font and spacing for document."""
        style = self.doc.styles['Normal']
        
        # Font settings
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(12)
        
        # Paragraph settings
        paragraph_format = style.paragraph_format
        paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
        paragraph_format.space_before = Pt(0)
        paragraph_format.space_after = Pt(0)
    
    def _set_margins(self):
        """Set 1-inch margins on all sides."""
        for section in self.doc.sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
    
    def _add_page_numbers(self):
        """Add page numbers to top right of every page."""
        section = self.doc.sections[0]
        header = section.header
        header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Add page number field
        run = header_para.add_run()
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = "PAGE"
        
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        
        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)
    
    def _create_title_page(self, title_data):
        """Create APA-formatted title page."""
        # Three blank lines at top
        for _ in range(3):
            self.doc.add_paragraph()
        
        # Title (bold, centered)
        title_para = self.doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.add_run(title_data['title'])
        title_run.bold = True
        
        # Blank line after title
        self.doc.add_paragraph()
        
        # Author, institution, course, instructor, date (all centered)
        for key in ['author', 'institution', 'course', 'instructor', 'date']:
            if key in title_data and title_data[key]:
                para = self.doc.add_paragraph(title_data[key])
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    def _create_body(self, body_content):
        """Create document body with proper heading levels and paragraphs."""
        i = 0
        while i < len(body_content):
            block = body_content[i]
            
            if block['type'] == 'heading':
                level = block['level']
                text = block['text']
                
                # Check for run-in headings (levels 4-5)
                next_block = body_content[i + 1] if i + 1 < len(body_content) else None
                
                if level in [4, 5] and next_block and next_block['type'] == 'paragraph':
                    # Create run-in heading
                    para = self.doc.add_paragraph()
                    para.paragraph_format.left_indent = Inches(0.5)
                    para.paragraph_format.first_line_indent = Inches(0)
                    
                    # Heading text (bold, italic for level 5)
                    heading_run = para.add_run(text + '. ')
                    heading_run.bold = True
                    if level == 5:
                        heading_run.italic = True
                    
                    # Paragraph text continues on same line
                    para.add_run(next_block['text'])
                    i += 2  # Skip next block since we processed it
                    continue
                else:
                    # Regular headings (levels 1-3 or standalone 4-5)
                    para = self.doc.add_paragraph(text)
                    
                    if level == 1:
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        para.runs[0].bold = True
                    elif level == 2:
                        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        para.runs[0].bold = True
                    elif level == 3:
                        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        para.runs[0].bold = True
                        para.runs[0].italic = True
                    elif level in [4, 5]:
                        # Standalone level 4/5 heading (fallback)
                        para.paragraph_format.left_indent = Inches(0.5)
                        heading_run = para.runs[0]
                        heading_run.bold = True
                        if level == 5:
                            heading_run.italic = True
                        # Add period if not already present
                        if not text.endswith('.'):
                            para.add_run('.')
            
            elif block['type'] == 'paragraph':
                para = self.doc.add_paragraph(block['text'])
                para.paragraph_format.first_line_indent = Inches(0.5)
            
            elif block['type'] == 'math':
                para = self.doc.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                formula_run = para.add_run(block['formula'])
                formula_run.italic = True
            
            i += 1
    
    def _create_references(self, references):
        """Create APA-formatted references section."""
        if not references:
            return
            
        # References heading (Level 1)
        heading = self.doc.add_paragraph('References')
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        heading.runs[0].bold = True
        
        # Reference entries with hanging indent
        for ref in references:
            if ref.strip():  # Only add non-empty references
                para = self.doc.add_paragraph(ref.strip())
                para.paragraph_format.left_indent = Inches(0.5)
                para.paragraph_format.first_line_indent = Inches(-0.5)


class TextParser:
    """Parse markdown-style text into structured content."""
    
    @staticmethod
    def parse(raw_text):
        """
        Parse markdown-style text into structured content blocks.
        
        Parameters:
        -----------
        raw_text : str
            Raw text with markdown headings, paragraphs, and references
        
        Returns:
        --------
        tuple: (body_content list, references list)
        """
        body_content = []
        references = []
        lines = raw_text.split('\n')
        
        in_references = False
        current_paragraph = []
        
        for i, line in enumerate(lines):
            stripped = line.strip()
            
            # Check for references section marker
            if stripped.lower() == 'references' or (stripped == '---' and i + 1 < len(lines)):
                # Check if next non-empty line after --- looks like a reference
                if stripped == '---':
                    for future_line in lines[i+1:]:
                        if future_line.strip():
                            # Simple heuristic: references often start with author names (capital letter)
                            if future_line.strip()[0].isupper():
                                in_references = True
                            break
                else:
                    in_references = True
                
                # Save any accumulated paragraph before entering references
                if current_paragraph:
                    body_content.append({
                        'type': 'paragraph',
                        'text': ' '.join(current_paragraph)
                    })
                    current_paragraph = []
                continue
            
            if in_references:
                if stripped:  # Non-empty reference line
                    references.append(stripped)
                continue
            
            # Parse headings (markdown style)
            if stripped.startswith('#'):
                # Save any accumulated paragraph
                if current_paragraph:
                    body_content.append({
                        'type': 'paragraph',
                        'text': ' '.join(current_paragraph)
                    })
                    current_paragraph = []
                
                # Count heading level
                level = 0
                for char in stripped:
                    if char == '#':
                        level += 1
                    else:
                        break
                
                text = stripped[level:].strip()
                body_content.append({
                    'type': 'heading',
                    'level': min(level, 5),
                    'text': text
                })
            
            # Parse math formulas (between double backticks or single backticks)
            elif '``' in stripped or '`' in stripped:
                if current_paragraph:
                    body_content.append({
                        'type': 'paragraph',
                        'text': ' '.join(current_paragraph)
                    })
                    current_paragraph = []
                
                # Extract formula (handle both `` and ` formats)
                formula_match = re.search(r'``(.+?)``', stripped)
                if not formula_match:
                    formula_match = re.search(r'`(.+?)`', stripped)
                
                if formula_match:
                    formula = formula_match.group(1)
                    body_content.append({
                        'type': 'math',
                        'formula': formula
                    })
            
            # Regular paragraph text
            elif stripped and not stripped.startswith('---'):
                current_paragraph.append(stripped)
            
            # Empty line ends paragraph
            elif current_paragraph:
                body_content.append({
                    'type': 'paragraph',
                    'text': ' '.join(current_paragraph)
                })
                current_paragraph = []
        
        # Add final paragraph if exists
        if current_paragraph:
            body_content.append({
                'type': 'paragraph',
                'text': ' '.join(current_paragraph)
            })
        
        return body_content, references


def create_apa_document(title_data, raw_text, output_path):
    """
    Main function to create an APA document from raw text.
    
    Parameters:
    -----------
    title_data : dict
        Title page information
    raw_text : str
        Raw document text with markdown formatting
    output_path : str
        Path to save the output document
    """
    # Parse the text
    parser = TextParser()
    body_content, references = parser.parse(raw_text)
    
    # Create the document
    formatter = APAFormatter()
    doc = formatter.create_document(title_data, body_content, references)
    
    # Save the document
    doc.save(output_path)
    print(f"✓ APA document created successfully: {output_path}")


# Example usage
if __name__ == '__main__':
    # Sample title data
    title_data = {
        'title': 'The Impact of Technology in Modern Education',
        'author': 'Jane Doe',
        'institution': 'University of Technology',
        'course': 'EDU-601: Advanced Learning Theories',
        'instructor': 'Dr. Alan Turing',
        'date': 'October 16, 2025'
    }
    
    # Sample raw text with markdown
    raw_text = """
# Introduction
Artificial intelligence is transforming the landscape of modern education, offering new possibilities for personalized learning and data-driven instruction.

## Background
The integration of technology in educational settings has evolved significantly over the past several decades. From the introduction of computers in classrooms to the current era of artificial intelligence and machine learning, each technological advancement has brought new opportunities and challenges.

### Early Developments
The 1980s marked the beginning of widespread computer adoption in schools. Initial applications focused on basic computer literacy and simple educational software.

#### Research on Early Computing
Smith (2015) conducted a comprehensive study on the effectiveness of early educational technology. The research demonstrated mixed results, with significant variations based on implementation quality and teacher training.

##### Specific Findings from Pilot Programs
Data from pilot programs in California schools showed a 15% improvement in mathematics scores. These early successes laid the groundwork for future technological integration.

The formula for calculating educational impact is: ``Impact = (Post_Score - Pre_Score) / Pre_Score × 100``

## Current Trends
Today's educational technology landscape is characterized by adaptive learning systems, learning analytics, and AI-powered tutoring systems.

---
References
American Psychological Association. (2020). Publication manual of the American Psychological Association (7th ed.). https://doi.org/10.1037/0000165-000
Johnson, S. (2023). AI in modern education. Healthcare Technology Review, 12(3), 45-67.
Smith, A., & Brown, B. (2015). Early educational computing: A retrospective analysis. Journal of Educational Technology, 28(2), 112-134.
Williams, T. (2022). Machine learning for personalized instruction. Academic Press.
"""
    
    # Create the document
    create_apa_document(
        title_data=title_data,
        raw_text=raw_text,
        output_path='/mnt/user-data/outputs/example_apa_paper.docx'
    )
